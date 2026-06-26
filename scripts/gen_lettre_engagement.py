"""
Generate lettre_engagement_cifre.docx from scratch using python-docx.
Run:  py scripts/gen_lettre_engagement.py
Output: presentation/lettre_engagement_cifre.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(BASE_DIR, "presentation", "lettre_engagement_cifre.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font  = style.font
font.name = "Calibri"
font.size = Pt(10)

def para(text="", bold=False, size=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
    return p

def add_run(p, text, bold=False, size=None):
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return run

# ── Header: sender block (left) ───────────────────────────────────────────────
p = para()
add_run(p, "Relations Humaines et Organisation", bold=True)

p = para()
add_run(p, "Gestion administrative des collaborateurs", bold=True)

para("TSA 70100")
para("75309 PARIS Cedex 09")

doc.add_paragraph()  # spacer

# ── Recipient block (right-aligned) ──────────────────────────────────────────
for line in ["ANRT", "SERVICE CIFRE", "33 RUE RENNEQUIN", "75017 PARIS"]:
    p = para(line, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph()

# ── Date ──────────────────────────────────────────────────────────────────────
para("Saint Denis, le 26 juin 2026")

doc.add_paragraph()

# ── Subject ───────────────────────────────────────────────────────────────────
p = para()
add_run(p, "Objet\u00a0: ", bold=True)
add_run(p, (
    "Engagement de Generali, dans le cadre du projet de th\u00e8se CIFRE par "
    "M Yassine Mannai, sous la direction de Caroline HILLAIRET (Laboratoire CREST, "
    "ENSAE Paris) et Anthony R\u00c9VEILLAC (INSA Toulouse, Institut de "
    "Math\u00e9matiques de Toulouse)."
))

doc.add_paragraph()

para("Madame, Monsieur,")

doc.add_paragraph()

# ── Body paragraphs ───────────────────────────────────────────────────────────
para(
    "Les soussign\u00e9es Sylvie PERETTI en sa qualit\u00e9 de Membre du Comit\u00e9 Ex\u00e9cutif en "
    "charge des Relations Humaines et Organisation et Virginie LILLE en sa qualit\u00e9 "
    "de Directeur Pilotage et Gestion RH de l\u2019Entreprise Generali France, dont le "
    "si\u00e8ge social est situ\u00e9 89 rue Taitbout, 75009 Paris,"
)

para(
    "s\u2019engagent \u00e0 accueillir M Yassine Mannai dans le cadre du projet de th\u00e8se CIFRE "
    "\u00ab\u00a0D\u00e9pendance factorielle conditionnelle aux r\u00e9gimes, proxy de capital supervis\u00e9 "
    "et couverture profonde pour le Total Portfolio Approach en assurance\u00a0\u00bb "
    "en collaboration avec le Laboratoire CREST et l\u2019Institut de Math\u00e9matiques de Toulouse."
)

para(
    "Ce recrutement s\u2019inscrit dans la logique de d\u00e9veloppement des capacit\u00e9s "
    "quantitatives et de recherche de l\u2019\u00e9quipe Allocation Strat\u00e9gique d\u2019Actifs (SAA) "
    "de Generali France. Cette \u00e9quipe, rattach\u00e9e \u00e0 la Direction des Investissements "
    "plac\u00e9e sous la responsabilit\u00e9 de M. Cedrik d\u2019Aviau de Ternay (Directeur des "
    "Investissements de Generali France), regroupe des expertises en finance "
    "quantitative, en gestion actif-passif (ALM), en mod\u00e9lisation actuarielle et "
    "en optimisation de portefeuille sous contraintes r\u00e9glementaires. Elle pilote "
    "l\u2019allocation strat\u00e9gique du bilan de Generali France dans le cadre de Solvabilit\u00e9\u00a0II, "
    "en s\u2019appuyant notamment sur l\u2019outil propri\u00e9taire du Groupe, GPS (Group Portfolio Solutions)."
)

para(
    "La gestion du risque de d\u00e9pendance entre facteurs \u00e9conomiques dans un contexte de "
    "changements de r\u00e9gime constitue un enjeu croissant pour les assureurs. L\u2019\u00e9pisode "
    "de 2022, au cours duquel la corr\u00e9lation entre actions et obligations s\u2019est invers\u00e9e "
    "simultan\u00e9ment sur tous les march\u00e9s, a r\u00e9v\u00e9l\u00e9 les limites des approches classiques "
    "d\u2019allocation par classes d\u2019actifs et motive le d\u00e9veloppement d\u2019une approche "
    "factorielle, dynamique et int\u00e9gr\u00e9e au sein de l\u2019\u00e9quipe SAA."
)

para(
    "Dans ce contexte, Generali, qui a pris connaissance des r\u00e8gles d\u2019attribution et de "
    "fonctionnement des CIFRE, s\u2019engage \u00e0 recruter M Yassine Mannai en contrat CDD de "
    "3 ans en l\u2019int\u00e9grant aux effectifs de l\u2019\u00e9quipe Allocation Strat\u00e9gique d\u2019Actifs. Il "
    "sera plac\u00e9 sous la responsabilit\u00e9 scientifique de M. Mamy RAMAMONJISOA, membre de "
    "l\u2019\u00e9quipe SAA. Generali mettra tout en \u0153uvre et donnera tous les moyens au "
    "doctorant pour que la CIFRE soit men\u00e9e \u00e0 terme dans les meilleures conditions."
)

doc.add_paragraph()

para("Nous vous prions d\u2019agr\u00e9er, Madame, Monsieur, l\u2019expression de nos salutations distingu\u00e9es.")

# spacer before signatures
for _ in range(3):
    doc.add_paragraph()

# ── Signature table ───────────────────────────────────────────────────────────
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
# remove borders
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

# Row 0: names
table.cell(0, 0).paragraphs[0].add_run("SYLVIE PERETTI").bold = True
table.cell(0, 1).paragraphs[0].add_run("VIRGINIE LILLE").bold = True

# Row 1: titles
table.cell(1, 0).paragraphs[0].add_run(
    "Membre du Comit\u00e9 Ex\u00e9cutif en charge des\nRelations Humaines et Organisation"
)
table.cell(1, 1).paragraphs[0].add_run("Directeur Pilotage et Gestion RH")

doc.add_paragraph()

# ── Footer note (small) ───────────────────────────────────────────────────────
footer_text = (
    "Generali Iard, Soci\u00e9t\u00e9 anonyme au capital de 94\u00a0630\u00a0300 euros \u2014 Entreprise r\u00e9gie par le Code des assurances \u2014 "
    "552\u00a0062\u00a0663 RCS Paris \u2014 IDU ADEME FR232327_01NBYI\n"
    "Generali Vie, Soci\u00e9t\u00e9 anonyme au capital de 341\u00a0059\u00a0488 euros \u2014 Fonds de Retraite Professionnelle Suppl\u00e9mentaire "
    "r\u00e9gi par le Code des assurances \u2014 880\u00a0265\u00a0418 RCS Paris \u2014 IDU ADEME FR232327_01NBYI\n"
    "Si\u00e8ge social\u00a0: 89 rue Taitbout \u2014 75009 Paris \u2014 Soci\u00e9t\u00e9s appartenant au Groupe Generali immatricul\u00e9 "
    "sur le registre italien des groupes d\u2019assurances sous le num\u00e9ro 026"
)
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run(footer_text)
run.font.size = Pt(6.5)

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
